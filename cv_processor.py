"""
CV Processor — Extract text from PDF and DOCX files.
"""

import io
from pathlib import Path
from typing import Optional
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from a PDF file."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def extract_cv_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from a CV file based on its extension.
    Supports PDF and DOCX formats.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")


def get_file_info(file_bytes: bytes, filename: str) -> dict:
    """Get basic file info."""
    ext = Path(filename).suffix.lower()
    size_mb = len(file_bytes) / (1024 * 1024)
    info = {
        "filename": filename,
        "format": ext.upper().replace(".", ""),
        "size_mb": round(size_mb, 2),
        "size_bytes": len(file_bytes),
    }
    # Count pages for PDF
    if ext == ".pdf":
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                info["pages"] = len(pdf.pages)
        except Exception:
            info["pages"] = "Unknown"
    elif ext in (".docx", ".doc"):
        try:
            doc = Document(io.BytesIO(file_bytes))
            info["paragraphs"] = len([p for p in doc.paragraphs if p.text.strip()])
        except Exception:
            info["paragraphs"] = "Unknown"
    return info


def validate_cv_file(file_bytes: bytes, filename: str, max_size_mb: int = 100) -> tuple[bool, str]:
    """
    Validate CV file format and size.
    Returns (is_valid, error_message).
    """
    ext = Path(filename).suffix.lower()
    if ext not in (".pdf", ".docx", ".doc"):
        return False, f"Format {ext} tidak didukung. Gunakan PDF atau DOCX."

    size_mb = len(file_bytes) / (1024 * 1024)
    if size_mb > max_size_mb:
        return False, f"Ukuran file {size_mb:.1f}MB melebihi batas {max_size_mb}MB."

    return True, ""
