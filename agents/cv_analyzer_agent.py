"""
CV Analyzer Agent — Reviews CV content and generates feedback.
Provides ATS score, improvement suggestions, and generates ATS-friendly CV.
Supports N8N webhook routing or direct OpenAI calls.
Export functions (DOCX/PDF) always run locally.
"""

import io
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib import colors
import config


REVIEW_PROMPT = """Kamu adalah CV Review Expert dengan pengalaman 10+ tahun di bidang HR dan recruitment. 
Tugasmu adalah menganalisis CV user dan memberikan feedback komprehensif.

Berikan output dalam format berikut (gunakan heading yang sama persis):

## 📊 ATS Score: [score]/100

## ✅ Kelebihan CV
- [point 1]
- [point 2]
...

## ⚠️ Area yang Perlu Diperbaiki
- [point 1]
- [point 2]
...

## 💡 Saran Perbaikan
- [saran spesifik 1]
- [saran spesifik 2]
...

## 🔑 Keywords yang Terdeteksi
[list keywords/skills yang ditemukan di CV]

## 📝 Ringkasan Profil
[ringkasan singkat profil kandidat berdasarkan CV]

Jawab dalam Bahasa Indonesia. Berikan feedback yang spesifik dan actionable."""


ATS_CV_PROMPT = """Kamu adalah CV Writer Expert. Berdasarkan CV asli user berikut, buat versi CV yang ATS-friendly dalam format Markdown yang terstruktur rapi.

Aturan Pembuatan & Desain:
1. **Nama & Informasi Kontak (Header)**:
   - Nama Lengkap wajib diletakkan di baris pertama dengan format `# [NAMA LENGKAP]`.
   - Informasi Kontak (Alamat, Telepon, Email, LinkedIn) di baris kedua. Pisahkan dengan simbol `|` atau `•`, misalnya: "Jakarta, Indonesia | +62 812-3456-7890 | email@example.com | linkedin.com/in/username"
2. **Bagian/Section Utama** (seperti PENDIDIKAN, PENGALAMAN KERJA, SKILLS, SERTIFIKASI):
   - Gunakan format heading level 2: `## [NAMA BAGIAN]` (semua huruf kapital).
   - Tepat di bawah setiap judul bagian ini (baris berikutnya), wajib tambahkan garis pemisah horizontal menggunakan `---` (tiga tanda hubung).
3. **Isi Pengalaman & Pendidikan**:
   - Tulis nama perusahaan/institusi dalam format **Tebal** (`**Nama Perusahaan**`) dan periode waktu di baris yang sama.
   - Pekerjaan/Gelar ditulis di bawahnya dengan format *Miring* (`*Gelar/Jabatan*`).
   - Gunakan bullet points (`-`) untuk setiap poin pencapaian atau tanggung jawab kerja.

Keluaran (Output) hanya berisi teks CV baru dalam format Markdown ini, tanpa ada kalimat pembuka (seperti "Berikut adalah...") atau penutup."""


def _build_target_job_context(target_job: dict) -> str:
    """Build target job context string for prompts."""
    if not target_job:
        return ""
    parts = []
    parts.append(f"\n\n--- POSISI YANG DITARGETKAN ---")
    parts.append(f"Posisi: {target_job.get('job_title', 'N/A')}")
    parts.append(f"Perusahaan: {target_job.get('company_name', 'N/A')}")
    desc = target_job.get('job_description', '')
    if desc and desc != 'N/A':
        parts.append(f"Deskripsi Posisi: {desc[:2000]}")
    parts.append("--- END ---")
    return "\n".join(parts)


def review_cv(cv_text: str, target_job: dict = None) -> dict:
    """
    Analyze CV and return structured feedback.
    If target_job is provided, feedback is tailored to that specific position.

    Returns dict with:
    - "feedback": AI-generated feedback markdown
    - "available": whether AI service is configured
    """
    # Build target context
    target_context = _build_target_job_context(target_job)

    # Enhance system prompt if target job is provided
    system_prompt = REVIEW_PROMPT
    if target_job:
        system_prompt += f"""\n\nIMPORTAN: User menargetkan posisi spesifik. 
Berikan feedback CV yang SPESIFIK dan TERARAH untuk posisi \"{target_job.get('job_title', '')}\". 
Analisis apakah CV user sudah cocok untuk posisi tersebut, 
identifikasi gap yang perlu diperbaiki, 
dan berikan saran konkret agar CV lebih menarik untuk posisi ini."""

    # Try N8N first
    if config.is_n8n_configured():
        try:
            from n8n_client import review_cv_n8n
            ai_text = review_cv_n8n(cv_text, target_job=target_job)
            if ai_text and not ai_text.startswith("N8N error") and not ai_text.startswith("Tidak dapat"):
                return {"feedback": ai_text, "available": True}
            else:
                return {"feedback": f"Error: {ai_text}", "available": True}
        except Exception as e:
            return {"feedback": f"Error N8N: {str(e)}", "available": True}

    # Fallback to local OpenAI
    if not config.is_openai_configured():
        return {
            "feedback": None,
            "available": False,
        }

    try:
        client = OpenAI(api_key=config.get_openai_api_key())
        user_content = f"Berikut adalah CV yang perlu di-review:\n\n{cv_text}{target_context}"
        response = client.chat.completions.create(
            model=config.OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.5,
            max_tokens=2500,
        )
        return {
            "feedback": response.choices[0].message.content,
            "available": True,
        }
    except Exception as e:
        return {
            "feedback": f"Error: {str(e)}",
            "available": True,
        }


def generate_ats_cv(cv_text: str, target_job: dict = None) -> dict:
    """
    Generate an ATS-friendly version of the CV.
    If target_job is provided, the ATS CV is optimized for that specific position.

    Returns dict with:
    - "ats_text": improved CV content as plain text
    - "available": whether AI service is configured
    """
    # Build target context
    target_context = _build_target_job_context(target_job)

    # Enhance system prompt if target job is provided
    system_prompt = ATS_CV_PROMPT
    if target_job:
        system_prompt += f"""\n\nIMPORTAN: Optimalkan CV ini KHUSUS untuk posisi \"{target_job.get('job_title', '')}\" di \"{target_job.get('company_name', '')}\".
Sesuaikan keywords, skills, dan pengalaman yang di-highlight agar relevan dengan posisi tersebut."""

    # Try N8N first
    if config.is_n8n_configured():
        try:
            from n8n_client import generate_ats_cv_n8n
            ai_text = generate_ats_cv_n8n(cv_text, target_job=target_job)
            if ai_text and not ai_text.startswith("N8N error") and not ai_text.startswith("Tidak dapat"):
                return {"ats_text": ai_text, "available": True}
            else:
                return {"ats_text": f"Error: {ai_text}", "available": True}
        except Exception as e:
            return {"ats_text": f"Error N8N: {str(e)}", "available": True}

    # Fallback to local OpenAI
    if not config.is_openai_configured():
        return {"ats_text": None, "available": False}

    try:
        client = OpenAI(api_key=config.get_openai_api_key())
        user_content = f"CV Asli:\n\n{cv_text}{target_context}"
        response = client.chat.completions.create(
            model=config.OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.4,
            max_tokens=3000,
        )
        return {
            "ats_text": response.choices[0].message.content,
            "available": True,
        }
    except Exception as e:
        return {"ats_text": f"Error: {str(e)}", "available": True}


# ─── Export Functions (always local, no N8N needed) ───────


def export_cv_to_docx(cv_text: str) -> bytes:
    """Export CV text to a formatted DOCX file. Returns bytes."""
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(33, 33, 33)

    def add_markdown_paragraph(text, paragraph_style=None, alignment=None, space_after=Pt(4), space_before=Pt(0)):
        if paragraph_style:
            p = doc.add_paragraph(style=paragraph_style)
        else:
            p = doc.add_paragraph()
        if alignment is not None:
            p.alignment = alignment
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = space_before
        
        # Split by bold markers
        import re
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                clean_part = part[2:-2]
                subparts = re.split(r'(\*.*?\*)', clean_part)
                for subpart in subparts:
                    if subpart.startswith('*') and subpart.endswith('*'):
                        r = p.add_run(subpart[1:-1])
                        r.bold = True
                        r.italic = True
                    else:
                        r = p.add_run(subpart)
                        r.bold = True
            else:
                subparts = re.split(r'(\*.*?\*)', part)
                for subpart in subparts:
                    if subpart.startswith('*') and subpart.endswith('*'):
                        r = p.add_run(subpart[1:-1])
                        r.italic = True
                    else:
                        r = p.add_run(subpart)
        return p

    # Parse and add content
    lines = cv_text.strip().split("\n")
    first_heading_added = False
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue

        # Handle header formatting
        if line.startswith("# ") and not first_heading_added:
            clean = line.lstrip("#").strip()
            if clean.startswith("**") and clean.endswith("**"):
                clean = clean[2:-2]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(clean)
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(17, 17, 17)
            first_heading_added = True
        elif first_heading_added and not line.startswith("## ") and not line == "---":
            # Contact info line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(15)
            # Remove any raw bold/italic marks for contact info
            clean_contact = line.replace("**", "").replace("*", "")
            run = p.add_run(clean_contact)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(85, 85, 85)
            first_heading_added = False
        elif line.startswith("## "):
            clean = line.lstrip("##").strip()
            if clean.startswith("**") and clean.endswith("**"):
                clean = clean[2:-2]
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean.upper())
            run.font.name = "Calibri"
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = RGBColor(17, 17, 17)
        elif line == "---":
            # Add horizontal line in Word using a paragraph border
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '8')  # thickness
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '111111')
                pBdr.append(bottom)
                pPr.append(pBdr)
            except Exception:
                p.add_run("____________________________________________________")
        elif line.startswith("- ") or line.startswith("• "):
            clean = line.lstrip("-•").strip()
            add_markdown_paragraph(clean, paragraph_style="List Bullet", space_after=Pt(3))
        else:
            add_markdown_paragraph(line, space_after=Pt(4))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_cv_to_pdf(cv_text: str) -> bytes:
    """Export CV text to a formatted PDF file. Returns bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=54,
        leftMargin=54,
        topMargin=45,
        bottomMargin=45,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    heading_style = ParagraphStyle(
        "CVHeading",
        parent=styles["Heading2"],
        fontSize=12,
        spaceAfter=2,
        spaceBefore=12,
        textColor=colors.HexColor("#111111"),
        fontName="Helvetica-Bold",
    )
    body_style = ParagraphStyle(
        "CVBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=4,
        leading=14,
        fontName="Helvetica",
    )
    bullet_style = ParagraphStyle(
        "CVBullet",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=3,
        leftIndent=15,
        leading=14,
        bulletIndent=5,
        fontName="Helvetica",
    )
    name_style = ParagraphStyle(
        "CVName",
        parent=styles["Heading1"],
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#111111"),
        fontName="Helvetica-Bold",
        spaceAfter=4,
    )
    contact_style = ParagraphStyle(
        "CVContact",
        parent=styles["Normal"],
        fontSize=9.5,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#444444"),
        spaceAfter=12,
        fontName="Helvetica",
    )

    import re
    def md_to_pdf_html(text: str) -> str:
        # Escape XML
        safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # Convert **bold** to <b>bold</b>
        safe = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', safe)
        # Convert *italic* to <i>italic</i>
        safe = re.sub(r'\*(.*?)\*', r'<i>\1</i>', safe)
        return safe

    elements = []
    lines = cv_text.strip().split("\n")
    first_heading_added = False

    for line in lines:
        line = line.strip()
        if not line:
            elements.append(Spacer(1, 4))
            continue

        if line.startswith("# ") and not first_heading_added:
            clean = line.lstrip("#").strip()
            if clean.startswith("**") and clean.endswith("**"):
                clean = clean[2:-2]
            elements.append(Paragraph(md_to_pdf_html(clean), name_style))
            first_heading_added = True
        elif first_heading_added and not line.startswith("## ") and not line == "---":
            # Strip bold/italic markup from contact lines
            clean_contact = line.replace("**", "").replace("*", "")
            elements.append(Paragraph(md_to_pdf_html(clean_contact), contact_style))
            first_heading_added = False
        elif line.startswith("## "):
            clean = line.lstrip("##").strip()
            if clean.startswith("**") and clean.endswith("**"):
                clean = clean[2:-2]
            elements.append(Paragraph(md_to_pdf_html(clean.upper()), heading_style))
        elif line == "---":
            elements.append(HRFlowable(
                width="100%", 
                thickness=1.5, 
                color=colors.HexColor("#111111"), 
                spaceBefore=1, 
                spaceAfter=8
            ))
        elif line.startswith("- ") or line.startswith("• "):
            clean = line.lstrip("-•").strip()
            elements.append(Paragraph(md_to_pdf_html(clean), bullet_style))
        else:
            elements.append(Paragraph(md_to_pdf_html(line), body_style))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()
